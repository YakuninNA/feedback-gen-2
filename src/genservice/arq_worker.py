import base64
import io
import json
from docx import Document
from arq.connections import logger
from config import redis_settings

from genservice.functionality.few_shot_prompting import (
    few_shot_categorization,
    few_shot_engineering_basics,
    few_shot_extraction,
    few_shot_experience,
    few_shot_polish,
    few_shot_requirements,
    few_shot_technical_skills,
    few_shot_candidate_name)
from genservice.functionality.langchain_pipeline import (
    parse_technical_requirements,
    run_qa_extraction_pipeline,
    run_ti_sections_extraction_pipeline,
    parse_candidate_name
)
from genservice.functionality.utility import general_requirements


async def process_transcript_task(
    ctx, json_file, requirements, position, filename, timestamp
):
    redis = ctx['redis']

    try:
        transcript_data = json.loads(json_file)

        feedback_name = await parse_candidate_name(
            few_shot_candidate_name=few_shot_candidate_name,
            transcript_name=filename,
            timestamp=timestamp
        )

        all_requirements = await parse_technical_requirements(
            requirements_text=requirements,
            general_requirements=general_requirements,
            few_shot_requirements=few_shot_requirements
        )

        extracted_data = await run_qa_extraction_pipeline(
            processed_data=transcript_data,
            tech_requirements=all_requirements['tech_requirements'],
            general_requirements=all_requirements['general_requirements'],
            categories=all_requirements['all_requirements'],
            few_shot_extraction=few_shot_extraction,
            few_shot_polish=few_shot_polish,
            few_shot_categorization=few_shot_categorization
        )

        feedback_components = await run_ti_sections_extraction_pipeline(
            soft_categorized_answers=extracted_data['soft_categorized_answers'],
            tech_categorized_answers=extracted_data['tech_categorized_answers'],
            position_name=position,
            tech_requirements=all_requirements['tech_requirements'],
            few_shot_experience=few_shot_experience,
            few_shot_engineering_basics=few_shot_engineering_basics,
            few_shot_technical_skills=few_shot_technical_skills
        )

        document = Document()
        document.add_heading(f"Feedback for {feedback_name}", level=1)
        document.add_paragraph("Generated Feedback:")
        for section, content in feedback_components.items():
            document.add_heading(section, level=2)
            document.add_paragraph(content if isinstance(content, str) else json.dumps(content, indent=4))

        file_like = io.BytesIO()
        document.save(file_like)
        file_like.seek(0)

        file_content = base64.b64encode(file_like.getvalue()).decode("utf-8")

        await redis.set(ctx['job_id'], json.dumps({"filename": f"{feedback_name}.docx", "file_content": file_content}), ex=3600)
        logger.info(f"Task completed successfully. Result stored for job ID {ctx['job_id']}")

    except Exception as e:
        logger.exception(f"An error occurred while processing the task: {e}")
        await redis.set(ctx['job_id'], json.dumps({"error": str(e)}), ex=3600)


class WorkerSettings:
    functions = [process_transcript_task]
    redis_settings = redis_settings
    max_jobs = 50
    poll_interval = 1
    timeout = 250
    job_timeout = 1000
    log_level = 'INFO'
